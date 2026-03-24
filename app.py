import streamlit as st
import pandas as pd
import time
import json
import plotly.express as px
from supabase import create_client, Client
from docx import Document
from io import BytesIO

# --- 1. DATABASE & DATA SETUP ---
URL = "https://tmbtnbxrrylulhgvnfjj.supabase.co"
KEY = "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJpc3MiOiJzdXBhYmFzZSIsInJlZiI6InRtYnRuYnhycnlsdWxoZ3ZuZmpqIiwicm9sZSI6ImFub24iLCJpYXQiOjE3NzQxMDQ2ODcsImV4cCI6MjA4OTY4MDY4N30.Fd1TPTCjN2u-_EOmkkqOb3TAKW8Q5RGv0AtAA85jW4s"
supabase: Client = create_client(URL, KEY)
SHEET_URL = "https://docs.google.com/spreadsheets/d/e/2PACX-1vSeTbSBHxYsciOesGpXt6ATZm_5aWVHQrS7tIFIaibmU4MZU-otPRsxUXG4egEP7P7jXdtL6CHhytAw/pub?output=csv"

@st.cache_data(ttl=1)
def load_data():
    try: 
        data = pd.read_csv(SHEET_URL)
        data.columns = [str(c).strip().lower() for c in data.columns]
        return data
    except: return pd.DataFrame()

df = load_data()

# --- 2. REPORT GENERATOR (Fixed for Special Characters) ---
def create_docx(name, score, total, script):
    doc = Document()
    doc.add_heading('VikidylEdu CBT - Official Result', 0)
    doc.add_paragraph(f"Student: {name}")
    doc.add_paragraph(f"Score: {score} / {total} ({int(score/total*100)}%)")
    for i, item in enumerate(script):
        doc.add_heading(f"Q{i+1}", level=2)
        status = "CORRECT" if item.get('ok') else "INCORRECT"
        doc.add_paragraph(f"Result: {status}")
        doc.add_paragraph(f"Correct Answer: {item.get('ca')}")
        doc.add_paragraph(f"Explanation: {item.get('ex', 'N/A')}")
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

def get_remark(score, total):
    pct = (score / total) * 100
    if pct >= 80: return "🌟 **Outstanding!** Excellent mastery. Keep it up!"
    elif pct >= 60: return "👍 **Good Job!** Solid performance, with minor gaps to fill."
    elif pct >= 45: return "📚 **Fair Effort.** You passed, but review the weak topics below."
    else: return "🛠️ **Revision Needed.** Please study the corrections and try again."

# --- 3. UI STYLE ---
st.set_page_config(page_title="VikidylEdu CBT", layout="wide")
st.markdown("""<style>
    .winner-box { background-color: #FFD700; padding: 10px; border-radius: 10px; color: #000; text-align: center; font-weight: bold; margin-bottom: 8px; }
    .report-card { background-color: #ffffff; padding: 15px; border-radius: 10px; border-left: 5px solid #1E3A8A; margin-bottom: 10px; border: 1px solid #e5e7eb; color: #000; }
    .weak-topic { color: #cf1322; font-weight: bold; background-color: #fff1f0; padding: 2px 5px; border-radius: 4px; }
</style>""", unsafe_allow_html=True)

# --- 4. LEADERBOARD SIDEBAR ---
with st.sidebar:
    st.title("VikidylEdu")
    try:
        res = supabase.table("leaderboard").select("name, score").order("score", desc=True).limit(3).execute()
        for i, entry in enumerate(res.data):
            n = entry['name'].split('|')[1].strip() if '|' in entry['name'] else entry['name']
            st.markdown(f"<div class='winner-box'>#{i+1} {n}: {entry['score']}</div>", unsafe_allow_html=True)
    except: pass
    role = st.selectbox("Portal", ["✍️ Student", "👨‍🏫 Teacher", "👨‍👩‍👧 Parent"])

# --- 5. STUDENT PORTAL ---
if role == "✍️ Student":
    if 'exam_active' not in st.session_state and 'final_score' not in st.session_state:
        st.header("📝 Registration")
        c1, c2 = st.columns(2)
        with c1: sch = st.text_input("School:")
        with c2: nm = st.text_input("Full Name:")
        
        cy, ce, cs = st.columns(3)
        with cy: 
            yrs = ["ALL YEARS"] + sorted(df['year'].unique().astype(str).tolist(), reverse=True) if not df.empty else ["2024"]
            yr = st.selectbox("Year", yrs)
        with ce: exm = st.selectbox("Exam", ["JAMB", "WAEC", "NECO", "BECE"])
        with cs: 
            # Fixes the TypeError by handling empty data
            sub_list = sorted(df['subject'].unique().tolist()) if not df.empty else ["No Data"]
            sub = st.selectbox("Subject", sub_list)
        
        # Slider only shows for ALL YEARS
        num_q = st.slider("Select Question Count", 5, 50, 20) if yr == "ALL YEARS" else 50
        
        if st.button("🚀 START EXAM"):
            if not sch or len(nm.strip().split()) < 2: st.error("Enter School and Full Name.")
            else:
                filt = (df['subject'].str.upper() == sub.upper()) & (df['exam'].str.upper() == exm.upper())
                if yr != "ALL YEARS": filt &= (df['year'].astype(str) == yr)
                q_df = df[filt]
                if not q_df.empty:
                    limit = min(len(q_df), num_q)
                    st.session_state.quiz_data = q_df.sample(n=limit).reset_index(drop=True)
                    st.session_state.update({"exam_active": True, "start_time": time.time(), "current_q": 0, "user_answers": {}, "db_id": f"{sch} | {nm} | {sub} | {yr}"})
                    st.rerun()
                else: st.warning("No questions found.")

    elif 'exam_active' in st.session_state:
        # Timer and Questions
        rem = max(0, 1800 - int(time.time() - st.session_state.start_time))
        st.subheader(f"⏱️ {rem//60:02d}:{rem%60:02d}")
        q_df, curr = st.session_state.quiz_data, st.session_state.current_q
        row = q_df.iloc[curr]

        if 'passage' in row and pd.notna(row['passage']) and str(row['passage']).strip() != "":
            with st.expander("📖 Passage / Instructions", expanded=True):
                st.markdown(f"*{row['passage']}*")

        st.markdown(f"### Q{curr+1}: {row['question']}")
        st.session_state.user_answers[curr] = st.radio("Choose Answer:", [row['a'], row['b'], row['c'], row['d']], key=f"q_{curr}")
        
        col1, col2, col3 = st.columns([1,1,2])
        with col1: 
            if st.button("⬅️ Previous") and curr > 0: st.session_state.current_q -= 1; st.rerun()
        with col2:
            if st.button("Next ➡️") and curr < len(q_df)-1: st.session_state.current_q += 1; st.rerun()
        with col3:
            if st.button("🏁 FINISH"):
                score, script = 0, []
                for i, r in q_df.iterrows():
                    ans = st.session_state.user_answers.get(i, "None")
                    cor = str(r['correct_answer']).strip().upper()
                    ok = str(ans).strip().upper() == cor
                    if ok: score += 1
                    script.append({"q": r['question'], "ua": ans, "ca": cor, "ok": ok, "ex": r.get('explanation', 'N/A'), "topic": r.get('topic', 'General')})
                supabase.table("leaderboard").upsert({"name": st.session_state.db_id, "score": score, "script": json.dumps(script), "total_q": len(q_df)}, on_conflict="name").execute()
                st.session_state.update({"final_score": score, "final_script": script})
                del st.session_state['exam_active']; st.rerun()

    if 'final_score' in st.session_state:
        st.balloons()
        score, total = st.session_state.final_score, len(st.session_state.final_script)
        st.markdown(f"<h1 style='text-align: center;'>Your Score: {score} / {total}</h1>", unsafe_allow_html=True)
        st.info(get_remark(score, total))

        t1, t2, t3 = st.tabs(["📊 Performance Analysis", "🔍 Corrections & Explanations", "📥 Download Report"])
        
        with t1:
            st.subheader("Topic Strength Analysis")
            s_df = pd.DataFrame(st.session_state.final_script)
            weaks = s_df[s_df['ok'] == False]['topic'].unique()
            if len(weaks) > 0:
                st.warning("⚠️ **Review Needed:** You missed questions in these topics. Focus your study here:")
                for w in weaks: st.markdown(f"- <span class='weak-topic'>{w}</span>", unsafe_allow_html=True)
            else: st.success("🌟 Perfect! You mastered all topics in this session.")
            st.plotly_chart(px.pie(values=[score, total-score], names=['Correct', 'Incorrect'], hole=0.5, color_discrete_sequence=['#28a745', '#dc3545']))

        with t2:
            st.subheader("Detailed Corrections")
            for i, item in enumerate(st.session_state.final_script):
                with st.expander(f"{'✅' if item['ok'] else '❌'} Question {i+1}"):
                    st.write(f"**Q:** {item['q']}")
                    st.write(f"**Correct Answer:** {item['ca']}")
                    st.info(f"💡 **Explanation:** {item['ex']}")

        with t3:
            docx = create_docx(st.session_state.db_id, score, total, st.session_state.final_script)
            st.download_button("📥 Get Word Report", data=docx, file_name="Result.docx")
            if st.button("Take Another Test"): st.session_state.clear(); st.rerun()

# --- 6. TEACHER/PARENT PORTALS ---
elif role == "👨‍🏫 Teacher":
    if st.text_input("Access Key:", type="password") == "Lagos2026":
        res = supabase.table("leaderboard").select("*").execute()
        if res.data:
            ld = pd.DataFrame(res.data)
            ld['Subject'] = ld['name'].apply(lambda x: x.split('|')[2].strip() if '|' in x else "N/A")
            st.plotly_chart(px.bar(ld.groupby('Subject')['score'].mean().reset_index(), x='Subject', y='score'))
            sel = st.selectbox("Student Audit:", ["-- Select --"] + ld['name'].tolist())
            if sel != "-- Select --":
                row = ld[ld['name'] == sel].iloc[0]
                scr = json.loads(row['script'])
                st.download_button("📥 Download Report", data=create_docx(row['name'], row['score'], len(scr), scr), file_name="Student_Report.docx")

elif role == "👨‍👩‍👧 Parent":
    s_in, n_in = st.text_input("School:"), st.text_input("Child Name:")
    if st.button("Find Result") and s_in and n_in:
        res = supabase.table("leaderboard").select("*").execute()
        match = [r for r in res.data if s_in.lower() in r['name'].lower() and n_in.lower() in r['name'].lower()]
        if match:
            for m in match:
                st.write(f"**{m['name'].split('|')[2]}**: {m['score']} Marks")
                st.download_button(f"📥 Report for {m['name'].split('|')[1]}", data=create_docx(m['name'], m['score'], 20, json.loads(m['script'])), file_name="Parent_Report.docx")
        else: st.error("No record found.")
